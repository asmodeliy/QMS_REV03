#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Spec-Center Document Parser and Indexer

Automatically extracts content from PDF, DOCX, XLSX files in spec_center uploads
and indexes them into RAG for LLM usage.
"""

import os
import json
from pathlib import Path
from typing import Dict, List, Optional, Any

try:
    from pypdf import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    try:
        from PyPDF2 import PdfReader
        PDF_AVAILABLE = True
    except ImportError:
        PDF_AVAILABLE = False
        print("Warning: pypdf/PyPDF2 not installed. PDF support limited.")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. DOCX support limited.")

try:
    import openpyxl
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False
    print("Warning: openpyxl not installed. XLSX support limited.")


class SpecCenterParser:
    """Parse and extract content from Spec-Center documents"""
    
    def __init__(self, spec_center_path: str = "uploads/spec_center"):
        self.spec_center_path = Path(spec_center_path)
        self.documents: Dict[str, Dict[str, Any]] = {}
    
    def extract_pdf_content(self, file_path: Path) -> Optional[str]:
        """Extract text from PDF file"""
        if not PDF_AVAILABLE:
            return None
        
        try:
            content = []
            with open(file_path, 'rb') as f:
                reader = PdfReader(f)
                for page_num, page in enumerate(reader.pages[:10]):  # Limit to first 10 pages
                    text = page.extract_text()
                    if text.strip():
                        content.append(f"[Page {page_num + 1}]\n{text}")
            
            return "\n\n".join(content) if content else None
        except Exception as e:
            print(f"Error extracting PDF {file_path.name}: {e}")
            return None
    
    def extract_docx_content(self, file_path: Path) -> Optional[str]:
        """Extract text from DOCX file"""
        if not DOCX_AVAILABLE:
            return None
        
        try:
            doc = Document(file_path)
            content = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text)
            
            for table in doc.tables:
                content.append("[TABLE]")
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    content.append(" | ".join(row_data))
            
            return "\n".join(content) if content else None
        except Exception as e:
            print(f"Error extracting DOCX {file_path.name}: {e}")
            return None
    
    def extract_xlsx_content(self, file_path: Path) -> Optional[str]:
        """Extract content from XLSX file"""
        if not XLSX_AVAILABLE:
            return None
        
        try:
            wb = openpyxl.load_workbook(file_path)
            content = []
            
            for sheet_name in wb.sheetnames[:3]:  # Limit to first 3 sheets
                ws = wb[sheet_name]
                content.append(f"[Sheet: {sheet_name}]")
                
                for row in ws.iter_rows(values_only=True):
                    if any(cell for cell in row if cell):
                        content.append(" | ".join(str(cell or "") for cell in row))
            
            return "\n".join(content) if content else None
        except Exception as e:
            print(f"Error extracting XLSX {file_path.name}: {e}")
            return None
    
    def extract_content(self, file_path: Path) -> Optional[str]:
        """Extract content from any supported file format"""
        if not file_path.exists():
            return None
        
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            return self.extract_pdf_content(file_path)
        elif suffix == '.docx':
            return self.extract_docx_content(file_path)
        elif suffix in ['.xlsx', '.xls']:
            return self.extract_xlsx_content(file_path)
        elif suffix in ['.txt', '.md']:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except Exception:
                return None
        
        return None
    
    def parse_all_documents(self) -> Dict[str, Dict[str, Any]]:
        """Parse all documents in spec_center directory"""
        self.documents = {}
        
        if not self.spec_center_path.exists():
            print(f"Warning: Spec center path not found: {self.spec_center_path}")
            return self.documents
        
        # Group files by base name (ignoring timestamps)
        file_groups: Dict[str, List[Path]] = {}
        
        for file_path in self.spec_center_path.iterdir():
            if file_path.is_file():
                # Extract base name (remove timestamp prefix)
                # Format: YYYYMMDD_HHMMSS_filename.ext
                parts = file_path.name.split('_', 2)
                if len(parts) >= 3:
                    base_name = parts[2]
                else:
                    base_name = file_path.name
                
                if base_name not in file_groups:
                    file_groups[base_name] = []
                file_groups[base_name].append(file_path)
        
        print(f"Found {len(file_groups)} document groups in spec_center")
        
        # Process each document group
        for base_name, file_list in file_groups.items():
            # Prioritize by format: PDF > DOCX > XLSX > TXT
            preferred_file = None
            for priority_ext in ['.pdf', '.docx', '.xlsx', '.txt']:
                for f in file_list:
                    if f.suffix.lower() == priority_ext:
                        preferred_file = f
                        break
                if preferred_file:
                    break
            
            if not preferred_file:
                preferred_file = file_list[0]
            
            print(f"Processing: {base_name}...")
            content = self.extract_content(preferred_file)
            
            if content:
                self.documents[base_name] = {
                    "file_path": str(preferred_file),
                    "file_name": base_name,
                    "file_size": preferred_file.stat().st_size,
                    "content": content[:5000],  # Limit content size
                    "content_preview": content[:200] + "..." if len(content) > 200 else content,
                    "full_content": content,
                    "keywords": self._extract_keywords(base_name, content)
                }
                print(f"  [OK] Extracted {len(content)} characters")
            else:
                print(f"  [FAIL] Could not extract content")
        
        return self.documents
    
    def _extract_keywords(self, filename: str, content: str) -> List[str]:
        """Extract keywords from filename and content"""
        keywords = []
        
        # From filename
        if "ISO" in filename.upper():
            keywords.append("ISO")
        if "26262" in filename:
            keywords.append("Functional Safety")
            keywords.append("ISO26262")
        if "AEC" in filename.upper():
            keywords.append("AEC-Q100")
            keywords.append("Automotive")
        if "RS-COP" in filename:
            keywords.append("Process")
            keywords.append("Development")
        if "RS-QM" in filename:
            keywords.append("Quality")
        if "RS-EM" in filename:
            keywords.append("Environment")
        
        # From content (first 1000 chars)
        content_lower = content[:1000].lower()
        if "safety" in content_lower:
            keywords.append("Safety")
        if "process" in content_lower:
            keywords.append("Process")
        if "quality" in content_lower:
            keywords.append("Quality")
        
        return list(set(keywords))
    
    def get_document_by_keyword(self, keyword: str) -> List[Dict[str, Any]]:
        """Get documents matching a keyword"""
        results = []
        keyword_lower = keyword.lower()
        
        for doc_name, doc_info in self.documents.items():
            # Check filename
            if keyword_lower in doc_name.lower():
                results.append(doc_info)
                continue
            
            # Check keywords
            if any(keyword_lower in k.lower() for k in doc_info.get("keywords", [])):
                results.append(doc_info)
                continue
            
            # Check content preview
            if keyword_lower in doc_info.get("content_preview", "").lower():
                results.append(doc_info)
        
        return results[:5]  # Max 5 results
    
    def find_spec_file(self, keyword: str) -> Optional[Dict[str, Any]]:
        """Find spec file by keyword and return file info with path
        
        Returns: 
        {
            'file_name': 'ISO-26262-1-2018.pdf',
            'file_path': 'uploads/spec_center/[timestamped_file]',
            'real_path': Path object to actual file,
            'keywords': ['ISO26262', 'Functional Safety'],
            'preview': 'First 300 chars of content',
            'content': 'Full content or excerpt'
        }
        """
        # Load documents if not already loaded
        if not self.documents:
            if not self.load_index(".spec_center_index.json"):
                self.parse_all_documents()
        
        keyword_lower = keyword.lower()
        
        # Search in documents
        for doc_name, doc_info in self.documents.items():
            # Strong match: keyword in filename
            if keyword_lower in doc_name.lower():
                return {
                    **doc_info,
                    'search_keyword': keyword,
                    'match_type': 'filename'
                }
            
            # Medium match: keyword in keywords
            if any(keyword_lower in k.lower() for k in doc_info.get("keywords", [])):
                return {
                    **doc_info,
                    'search_keyword': keyword,
                    'match_type': 'keyword'
                }
            
            # Weak match: keyword in content
            if keyword_lower in doc_info.get("content_preview", "").lower():
                return {
                    **doc_info,
                    'search_keyword': keyword,
                    'match_type': 'content'
                }
        
        return None
    
    def find_spec_files(self, keyword: str, limit: int = 10) -> List[Dict[str, Any]]:
        """Find spec files matching keyword - returns multiple results
        
        Returns list of matching files with web paths for linking
        """
        # Load documents if not already loaded
        if not self.documents:
            if not self.load_index(".spec_center_index.json"):
                self.parse_all_documents()
        
        keyword_lower = keyword.lower()
        results = []
        
        # Search in order of relevance
        for doc_name, doc_info in self.documents.items():
            match_type = None
            
            # Strong match: keyword in filename
            if keyword_lower in doc_name.lower():
                match_type = 'filename'
            # Medium match: keyword in keywords
            elif any(keyword_lower in k.lower() for k in doc_info.get("keywords", [])):
                match_type = 'keyword'
            # Weak match: keyword in content
            elif keyword_lower in doc_info.get("content_preview", "").lower():
                match_type = 'content'
            
            if match_type:
                # Find actual file in uploads/spec_center
                actual_file = self._find_actual_file(doc_name)
                
                results.append({
                    **doc_info,
                    'search_keyword': keyword,
                    'match_type': match_type,
                    'actual_file_name': actual_file.name if actual_file else doc_name,
                    'web_path': f"uploads/spec_center/{actual_file.name}" if actual_file else f"uploads/spec_center/{doc_name}"
                })
        
        return results[:limit]
    
    def _find_actual_file(self, doc_name: str) -> Optional[Path]:
        """Find actual file in uploads/spec_center with timestamp prefix"""
        if not self.spec_center_path.exists():
            return None
        
        # Search for file matching doc_name
        for file_path in self.spec_center_path.iterdir():
            if file_path.is_file():
                # Extract base name (remove timestamp)
                parts = file_path.name.split('_', 2)
                if len(parts) >= 3:
                    base_name = parts[2]
                else:
                    base_name = file_path.name
                
                if base_name == doc_name:
                    return file_path
        
        return None
    
    def save_index(self, index_path: str = ".spec_center_index.json"):
        """Save parsed documents to JSON index"""
        try:
            index_data = {}
            for doc_name, doc_info in self.documents.items():
                # Don't include full content in index (too large)
                index_data[doc_name] = {
                    "file_path": doc_info["file_path"],
                    "file_name": doc_info["file_name"],
                    "file_size": doc_info["file_size"],
                    "content_preview": doc_info["content_preview"],
                    "keywords": doc_info["keywords"]
                }
            
            with open(index_path, 'w', encoding='utf-8') as f:
                json.dump(index_data, f, indent=2, ensure_ascii=False)
            
            print(f"Index saved to {index_path}")
        except Exception as e:
            print(f"Error saving index: {e}")
    
    def load_index(self, index_path: str = ".spec_center_index.json") -> bool:
        """Load parsed documents from JSON index"""
        if not os.path.exists(index_path):
            return False
        
        try:
            with open(index_path, 'r', encoding='utf-8') as f:
                index_data = json.load(f)
            
            # Load full content for each document
            for doc_name, doc_info in index_data.items():
                file_path = Path(doc_info["file_path"])
                content = self.extract_content(file_path)
                
                self.documents[doc_name] = {
                    **doc_info,
                    "full_content": content or ""
                }
            
            print(f"Loaded {len(self.documents)} documents from index")
            return True
        except Exception as e:
            print(f"Error loading index: {e}")
            return False


def index_spec_center_to_rag(rag_indexer=None):
    """Parse spec-center documents and add to RAG"""
    parser = SpecCenterParser(spec_center_path="uploads/spec_center")
    documents = parser.parse_all_documents()
    
    if not documents:
        return
    
    # Index to RAG if available
    if rag_indexer is None:
        try:
            from modules.mcp.rag_indexer import RAGIndexer
            rag_indexer = RAGIndexer(db_path="rag_knowledge_base.db")
        except Exception:
            print("RAG indexer not available")
            return
    
    print(f"\nIndexing {len(documents)} documents to RAG...")
    indexed_count = 0
    
    for doc_name, doc_info in documents.items():
        try:
            # Create structured document content
            content = f"""# {doc_name}

## Keywords
{', '.join(doc_info['keywords'])}

## Content
{doc_info['full_content']}"""
            
            rag_indexer.add_document(
                file_path=doc_info["file_name"],
                content=content,
                file_type="specification"
            )
            indexed_count += 1
            print(f"  ✓ Indexed: {doc_name}")
        except Exception as e:
            print(f"  ✗ Failed to index {doc_name}: {e}")
    
    print(f"Successfully indexed {indexed_count}/{len(documents)} documents")


if __name__ == "__main__":
    # Test
    parser = SpecCenterParser()
    documents = parser.parse_all_documents()
    parser.save_index()
    
    print("\n" + "=" * 60)
    print(f"Successfully parsed {len(documents)} documents")
    print("=" * 60)
